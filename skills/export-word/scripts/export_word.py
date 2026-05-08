#!/usr/bin/env python3
"""
专业 Word 数据分析报告导出 — WebSQL AI 平台 v2.0
模块化架构：共享基础设施 + 模板系统 + 构建器模式
生成符合中国商务标准的专业排版数据分析报告。
"""

import json
import os
import sys

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_SKILL_DIR = os.path.dirname(_SCRIPT_DIR)
_SKILLS_DIR = os.path.dirname(_SKILL_DIR)
sys.path.insert(0, _SCRIPT_DIR)
sys.path.insert(0, _SKILLS_DIR)

from shared.config import SkillConfig
from shared.colors import ColorPalette
from shared.logger import SkillLogger
from shared.utils import (
    clean_surrogates, format_date_cn, generate_report_id, ensure_output_dir, safe_json_dumps, strip_markdown
)
from shared.exceptions import SkillError, ValidationError, FileGenerationError

from word_templates.chinese_report import ChineseReportTemplate

from word_builders.cover import CoverBuilder
from word_builders.summary import SummaryBuilder
from word_builders.overview import OverviewBuilder
from word_builders.statistics import StatisticsBuilder
from word_builders.visualization import VisualizationBuilder
from word_builders.findings import FindingsBuilder
from word_builders.appendix import AppendixBuilder
from word_builders.header_footer import HeaderFooterBuilder


logger = SkillLogger("websql.skill.export_word")


class WordExporter:
    """Word 报告导出器"""

    def __init__(self, template_name="chinese_report", config=None):
        self.config = config or SkillConfig()
        scheme = ColorPalette.get("chinese_business")
        self.template = ChineseReportTemplate(scheme=scheme)
        self._init_builders()

    def _init_builders(self):
        self.cover = CoverBuilder(self.template)
        self.summary = SummaryBuilder(self.template)
        self.overview = OverviewBuilder(self.template)
        self.statistics = StatisticsBuilder(self.template)
        self.visualization = VisualizationBuilder(self.template)
        self.findings = FindingsBuilder(self.template)
        self.appendix = AppendixBuilder(self.template)
        self.header_footer = HeaderFooterBuilder(self.template)

    def _validate(self, d):
        if not d.get("title"):
            raise ValidationError("缺少必要字段: title")
        if "data" in d and not isinstance(d["data"], list):
            raise ValidationError("data 字段必须是数组")

    def export(self, d):
        self._validate(d)
        doc = self.template.create_document()

        title = d.get("title", "数据分析报告")
        subtitle = d.get("subtitle", "专业数据分析报告")
        dept = d.get("dept", "")
        columns = d.get("columns", [])
        data_rows = d.get("data", [])
        chart_paths = d.get("chartPaths", [])
        findings_list = d.get("findings", [])
        output_path = d.get("outputPath", "exports/report.docx")
        include_charts = d.get("includeCharts", True)
        numeric_cols = d.get("numericColumns", [])
        numeric_stats = d.get("numericStats", [])
        mode = d.get("mode", "data")

        logger.info(f"标题: {title}, 记录数: {len(data_rows)}, 图表数: {len(chart_paths)}, 模式: {mode}")

        if mode == "content":
            self._export_content_mode(doc, d, title, output_path)
        else:
            self.cover.build(doc, title, subtitle, dept)
            self.summary.build(doc, len(data_rows), len(columns), numeric_cols, numeric_stats)
            self.overview.build(doc, columns, data_rows)
            self.statistics.build(doc, numeric_stats, data_rows)

            if include_charts and chart_paths:
                self.visualization.build(doc, chart_paths)

            self.findings.build(doc, findings_list)
            self.appendix.build(doc, columns, data_rows)

            self.header_footer.build(doc, title)

        ensure_output_dir(output_path)
        doc.save(output_path)

        result = {
            "success": True,
            "path": output_path,
            "recordCount": len(data_rows),
            "message": f"Word 报告已生成，共 {len(data_rows)} 条记录",
        }
        logger.info(f"[export_word] 完成: {output_path}")
        print(json.dumps(result, ensure_ascii=False))
        return output_path

    def _export_content_mode(self, doc, d, title, output_path):
        sections = d.get("sections", [])
        self.cover.build(doc, title, d.get("subtitle", ""), d.get("dept", ""))

        for i, sec in enumerate(sections):
            self.template._heading(doc, 1, i + 2, sec.get("title", f"第{i+2}章"))
            blocks = sec.get("blocks") or []

            for b in blocks:
                bt = b.get("type", "")
                bc = strip_markdown(b.get("content", ""))

                if bt == "heading1":
                    self.template._heading(doc, 1, b.get("num"), bc)
                elif bt == "heading2":
                    self.template._heading(doc, 2, None, bc)
                elif bt == "heading3":
                    self.template._heading(doc, 3, None, bc)
                elif bt == "paragraph":
                    self.template._body_para(doc, bc)
                elif bt == "list":
                    for item in bc.split("\n"):
                        clean = item.strip("-  ").strip()
                        if clean:
                            p = doc.add_paragraph()
                            self.template._para_format(p, before=1, after=1, spacing=1.4,
                                                       indent_first=self.config.get("layout", "word", "first_indent_cm"))
                            self.template._run(p, f"•  {clean_surrogates(clean)}", size=11,
                                               color=self.template.scheme["text_primary"])
                elif bt == "table":
                    lines = [l.strip() for l in bc.strip().split("\n") if l.strip()]
                    if len(lines) >= 2:
                        headers = [h.strip() for h in lines[0].split("|") if h.strip()]
                        rows = [[v.strip() for v in line.split("|") if v.strip()] for line in lines[2:]]
                        if rows:
                            self.template._create_table(doc, headers, rows)
                elif bt == "image":
                    path = b.get("path", "")
                    caption = b.get("caption", "")
                    if path and os.path.exists(path):
                        from docx.shared import Inches
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        if caption:
                            p = doc.add_paragraph()
                            self.template._para_format(p, before=10, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
                            self.template._run(p, caption, size=10, bold=True,
                                               color=self.template.scheme["primary"])
                        doc.add_picture(path, width=Inches(5.6))

        self.header_footer.build(doc, title)


def main():
    logger.info("[export_word] 开始生成 Word 报告 v2.0")
    sys.stdin.reconfigure(encoding='utf-8')
    sys.stdout.reconfigure(encoding='utf-8')
    try:
        raw = sys.stdin.read()
        d = json.loads(raw)
    except json.JSONDecodeError as e:
        logger.error(f"[export_word] JSON 解析失败: {e}")
        print(json.dumps({"success": False, "error": f"JSON解析失败: {e}"}))
        sys.exit(1)

    try:
        exporter = WordExporter()
        exporter.export(d)
    except SkillError as e:
        logger.error(f"[export_word] {e}")
        print(json.dumps(e.to_dict(), ensure_ascii=False))
        sys.exit(1)
    except Exception as e:
        logger.exception(f"[export_word] 未预期的错误")
        print(json.dumps({"success": False, "error": str(e)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
