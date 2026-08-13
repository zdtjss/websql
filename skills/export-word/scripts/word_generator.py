# -*- coding: utf-8 -*-
"""word_generator.py — 模板驱动的 Word 报告导出脚本（渲染层唯一引擎）。

版式定义在 skills/export-word/templates/report_template.docx（docxtpl 模板：
封面占位、目录域、样式集、正文骨架），本脚本组装渲染上下文 + 渲染 +
表格/图片后处理插入。
换肤/换版式：直接编辑模板文件（或改 skills/lib/build_templates.py 后重新生成）。

JSON 契约（与 SKILL.md 一致）：
    data 模式:    {mode, title, columns, data, numericColumns, numericStats,
                   findings, chartPaths, includeCharts, outputPath}
    content 模式: {mode, title, sections:[{title, level, blocks:[...]}]},
                   markdown?, outputPath}
"""

import json
import os
import re
import sys
import tempfile
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docxtpl import DocxTemplate

# lib 目录加入 sys.path（兼容从任意工作目录执行）
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "lib"))

from chart_common import LightChartStyle, create_chart  # noqa: E402
from md_blocks import markdown_to_sections, normalize_sections  # noqa: E402

_CHART_STYLE = LightChartStyle()

# 后处理标记：模板中的 table/image 块渲染为标记段落，渲染后替换为真实内容
TABLE_TAG_RE = re.compile(r'%%WEBSQL_TABLE:(\d+)%%')
IMAGE_TAG_RE = re.compile(r'%%WEBSQL_IMAGE:(\d+)%%')


class Theme:
    """Word 浅色主题配色（与模板样式集保持一致）。"""
    HEADING1 = RGBColor(0x1A, 0x23, 0x7E)
    HEADING2 = RGBColor(0x28, 0x35, 0x93)
    HEADING3 = RGBColor(0x00, 0x71, 0xB8)
    PRIMARY = RGBColor(0x00, 0x71, 0xB8)
    BODY = RGBColor(0x2C, 0x2C, 0x2C)
    GRAY = RGBColor(0x80, 0x80, 0x80)
    TABLE_HEADER_BG = "1A237E"
    TABLE_ROW_ALT = "F0F4FF"
    CHART_BG = '#FFFFFF'
    CHART_FACE = '#FFFFFF'
    CHART_GRID = '#D0D7E2'


def _set_ea(run, typeface='微软雅黑'):
    run.font.name = typeface
    run._element.rPr.rFonts.set(qn('w:eastAsia'), typeface)


# ═══════════════════════════════════════════════════════════════
# 渲染上下文组装
# ═══════════════════════════════════════════════════════════════

GROUP_COLUMNS = {"表名", "table_name", "TABLE_NAME", "表名称"}


def _fmt_num(v):
    """数值安全格式化：None/字符串等异常值不崩溃。"""
    try:
        return f"{float(v):.2f}"
    except (TypeError, ValueError):
        return str(v) if v is not None else "N/A"


def _is_number(s):
    t = s.replace(",", "").replace("%", "")
    try:
        float(t)
        return True
    except ValueError:
        return False


def _to_float(v):
    try:
        return float(str(v).replace(",", ""))
    except (TypeError, ValueError):
        return 0.0


def _blank_block(kind):
    return {"kind": kind, "content": "", "items": [], "lines": []}


def _image_block(index):
    return {"kind": "image", "content": "", "items": [], "lines": [],
            "index": index}


def _table_block(index):
    return {"kind": "table", "content": "", "items": [], "lines": [],
            "index": index}


def _section(title, level=1, blocks=None):
    return {"title": title, "level": level, "blocks": blocks or []}


def build_context_data(data, title):
    """data 模式：组装 docxtpl 上下文 + 表格/图片后处理数据。

    返回 (context, tables_map, images_map, temp_files)
    """
    columns = data.get("columns", [])
    rows = data.get("data", [])
    findings = data.get("findings", [])
    numeric_stats = data.get("numericStats", [])
    chart_paths = data.get("chartPaths", [])
    include_charts = data.get("includeCharts", False)

    tables_map = {}   # index -> (headers, rows)
    images_map = {}   # index -> (png_path, caption)
    temp_files = []

    sections = []

    # 数据概览
    overview_blocks = [_blank_block("paragraph")]
    overview_blocks[0]["content"] = (
        f"本次分析共返回 {len(rows)} 条记录，包含 {len(columns)} 个字段。")
    sections.append(_section("数据概览", 1, overview_blocks))

    # 关键指标
    if numeric_stats:
        blocks = []
        for ns in numeric_stats:
            col_name = ns.get("column", "")
            b = _blank_block("paragraph")
            b["content"] = (
                f"{col_name}: 计数={ns.get('count', 0)}, "
                f"均值={_fmt_num(ns.get('avg', 0))}, "
                f"最小={_fmt_num(ns.get('min', 0))}, "
                f"最大={_fmt_num(ns.get('max', 0))}, "
                f"标准差={_fmt_num(ns.get('stddev', 0))}")
            blocks.append(b)
        sections.append(_section("关键指标", 2, blocks))

    # 分析洞察
    if findings:
        sections.append(_section(
            "分析洞察", 2, [{"kind": "bullets", "content": "",
                            "items": [str(f) for f in findings], "lines": []}]))

    # 数据明细（分组或整体表格）
    if columns and rows:
        group_col = next((c for c in columns if c in GROUP_COLUMNS), None)
        detail_blocks = []
        if group_col:
            from collections import OrderedDict
            groups = OrderedDict()
            for row in rows:
                gk = str(row.get(group_col, ""))
                groups.setdefault(gk, []).append(row)
            detail_blocks.append({
                "kind": "paragraph",
                "content": (f"共 {len(rows)} 条记录，按「{group_col}」"
                            f"分为 {len(groups)} 组展示。"),
                "items": [], "lines": []})
            other_cols = [c for c in columns if c != group_col]
            for gk, gk_rows in groups.items():
                sub = _blank_block("subheading")
                sub["content"] = f"{group_col}: {gk}（{len(gk_rows)} 条）"
                detail_blocks.append(sub)
                tb = _table_block(len(tables_map))
                detail_blocks.append(tb)
                tables_map[tb["index"]] = (
                    other_cols,
                    [[str(row.get(c, "")) for c in other_cols] for row in gk_rows])
        else:
            tb = _table_block(len(tables_map))
            detail_blocks.append(tb)
            tables_map[tb["index"]] = (
                columns,
                [[str(row.get(c, "")) for c in columns] for row in rows])
        sections.append(_section("数据明细", 2, detail_blocks))

    # 图表：data 模式自动生成柱状图（补齐原"data 模式不自动生成图表"的缺口）
    chart_section_blocks = []
    auto_chart_made = False
    if columns and rows and len(columns) >= 2:
        y_col = next((c for c in columns[1:] if _is_number(str(rows[0].get(c, "")))),
                     None)
        if y_col and len(rows) > 1:
            try:
                fig = create_chart(_CHART_STYLE, "bar", {
                    "title": f"{y_col} 分布",
                    "categories": [str(r.get(columns[0], ""))[:12] for r in rows[:10]],
                    "series": [{"name": y_col, "values": [
                        _to_float(r.get(y_col, 0)) for r in rows[:10]]}],
                })
                png = os.path.join(tempfile.gettempdir(),
                                   f"word_chart_auto_{len(temp_files)}.png")
                fig.savefig(png, dpi=150, bbox_inches='tight',
                            facecolor='white', edgecolor='none')
                plt.close(fig)
                temp_files.append(png)
                img = _image_block(len(images_map))
                chart_section_blocks.append(img)
                images_map[img["index"]] = (png, f"{y_col} 分布")
                auto_chart_made = True
            except Exception:
                auto_chart_made = False

    # 外部图表 PNG（Go 薄封装不再生成，兼容 skill 路径的 chartPaths）
    if include_charts and chart_paths:
        for cp in chart_paths:
            if os.path.exists(cp):
                img = _image_block(len(images_map))
                chart_section_blocks.append(img)
                images_map[img["index"]] = (cp, "")
    if chart_section_blocks:
        sections.append(_section("图表分析", 2, chart_section_blocks))
    _ = auto_chart_made  # 记录行为供日志，暂不输出

    context = {
        "title": title, "subtitle": "数据分析报告",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "author": "", "org": "",
        "sections": sections,
    }
    return context, tables_map, images_map, temp_files


def build_context_content(data, title):
    """content 模式：组装 docxtpl 上下文。

    markdown 字段优先（原始 Markdown 文本），否则用 sections 旧契约。
    """
    if data.get("markdown"):
        sections = markdown_to_sections(data["markdown"])
    else:
        sections = normalize_sections(data.get("sections", []))

    tables_map = {}
    images_map = {}
    temp_files = []

    # 为 table/image block 分配 index，并生成图表 PNG
    out_sections = []
    for sec in sections:
        out_blocks = []
        for b in sec.get("blocks", []):
            kind = b.get("kind", "paragraph")
            nb = {"kind": kind, "content": b.get("content", ""),
                  "items": b.get("items", []), "lines": b.get("lines", [])}
            if kind == "table":
                headers = b.get("headers", [])
                rows = b.get("rows", [])
                if headers and rows:
                    nb["index"] = len(tables_map)
                    tables_map[nb["index"]] = (headers, rows)
                else:
                    nb = None  # 空表格丢弃
            elif kind == "image":
                labels = b.get("labels", [])
                values = b.get("values", [])
                if labels and values:
                    try:
                        fig = create_chart(_CHART_STYLE, b.get("chart_type", "bar"), {
                            "title": b.get("chart_title", ""),
                            "categories": labels,
                            "series": [{"name": b.get("chart_title", ""),
                                        "values": values}],
                        })
                        png = os.path.join(
                            tempfile.gettempdir(),
                            f"word_chart_{len(temp_files)}.png")
                        fig.savefig(png, dpi=150, bbox_inches='tight',
                                    facecolor='white', edgecolor='none')
                        plt.close(fig)
                        temp_files.append(png)
                        nb["index"] = len(images_map)
                        images_map[nb["index"]] = (png, b.get("chart_title", ""))
                    except Exception:
                        nb = _blank_block("paragraph")
                        nb["content"] = f"[图表渲染失败: {b.get('chart_title', '')}]"
                else:
                    nb = None
            if nb is not None:
                out_blocks.append(nb)
        out_sections.append({"title": sec.get("title", ""),
                             "level": sec.get("level", 2),
                             "blocks": out_blocks})

    context = {
        "title": title, "subtitle": "数据分析报告",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "author": "", "org": "",
        "sections": out_sections,
    }
    return context, tables_map, images_map, temp_files


# ═══════════════════════════════════════════════════════════════
# 后处理：标记段落 → 表格 / 图片
# ═══════════════════════════════════════════════════════════════

def _build_table(doc, headers, rows):
    """构造数据表格（表头底纹、斑马纹），返回已插入文档的 Table。"""
    n_cols = max(1, len(headers))
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = 1  # CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_ea(run)
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), Theme.TABLE_HEADER_BG)
        cell._element.get_or_add_tcPr().append(shd)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row[:n_cols]):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            _set_ea(run)
            if r_idx % 2 == 1:
                from docx.oxml import OxmlElement
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), Theme.TABLE_ROW_ALT)
                cell._element.get_or_add_tcPr().append(shd)
    return table


def _postprocess(doc, tables_map, images_map):
    """把标记段落替换为真实表格/图片。"""
    for para in list(doc.paragraphs):
        text = para.text.strip()

        m = TABLE_TAG_RE.fullmatch(text)
        if m:
            idx = int(m.group(1))
            headers, rows = tables_map.get(idx, ([], []))
            if headers:
                tbl = _build_table(doc, headers, rows)
                para._p.addprevious(tbl._tbl)
            para._p.getparent().remove(para._p)
            continue

        m = IMAGE_TAG_RE.fullmatch(text)
        if m:
            idx = int(m.group(1))
            img_path, caption = images_map.get(idx, (None, ""))
            if img_path and os.path.exists(img_path):
                for r in list(para.runs):
                    r.text = ""
                run = para.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cp = para.insert_paragraph_before()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(caption)
                    cr.font.size = Pt(9)
                    cr.font.color.rgb = Theme.GRAY
                    cr.font.italic = True
                    _set_ea(cr)
            else:
                para._p.getparent().remove(para._p)
            continue


def _cleanup_empty_tags(doc):
    """移除渲染后残留的空标签段落（Tiny 字号的 for/endfor 段）。"""
    for para in list(doc.paragraphs):
        if para.text.strip() in ("",) and para.style.name == "Normal":
            # 仅清理由标签段产生的空段：字体极小且无其他格式
            runs = para.runs
            if runs and all((r.font.size or Pt(11)).pt <= 1 for r in runs):
                para._p.getparent().remove(para._p)


# ═══════════════════════════════════════════════════════════════
# 入口
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    # Windows 下默认 stdin/stdout 编码可能为 GBK（cp936），显式重配置为 UTF-8
    try:
        sys.stdin.reconfigure(encoding='utf-8')
        sys.stdout.reconfigure(encoding='utf-8')
    except (AttributeError, ValueError):
        pass

    try:
        raw = sys.stdin.read()
        # execute 工具不支持直接向子进程 stdin 写入，Agent 可能通过
        # `python xxx.py < input.json` 重定向或位置参数传入 JSON 文件
        if not raw and len(sys.argv) > 1:
            try:
                with open(sys.argv[1], "r", encoding="utf-8-sig") as f:
                    raw = f.read()
            except OSError as e:
                print(json.dumps({"success": False,
                                  "error": f"读取输入文件失败: {e}"},
                                 ensure_ascii=False))
                sys.exit(1)
        if not raw:
            print(json.dumps({"success": False, "error": "stdin 为空"},
                             ensure_ascii=False))
            sys.exit(1)
        data = json.loads(raw)
    except Exception as e:
        print(json.dumps({"success": False, "error": f"解析输入失败: {e}"},
                         ensure_ascii=False))
        sys.exit(1)

    temp_files = []
    try:
        mode = data.get("mode", "data")
        title = data.get("title", "数据分析报告")
        output_path = data.get("outputPath", "exports/output.docx")
        # 规范化输出路径：去掉前导 / 使其成为相对路径，确保 Windows 下
        # 保存到项目 exports/ 目录（/exports/x.docx 会解析到驱动器根）
        if output_path.startswith("/"):
            output_path = output_path.lstrip("/")

        template_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            "..", "templates", "report_template.docx")

        if mode == "content":
            context, tables_map, images_map, temp_files = \
                build_context_content(data, title)
        else:
            context, tables_map, images_map, temp_files = \
                build_context_data(data, title)

        tpl = DocxTemplate(template_path)
        tpl.render(context)

        # docxtpl 的 map_tree 把 lxml 裸元素替换进 python-docx 元素树，
        # 与 python-docx 1.2 的自定义元素类不兼容（访问 .paragraphs 会报
        # 'lxml._Element' has no attribute 'p_lst'）。改为落盘重载获得
        # 渲染后的 python-docx Document，与库内部实现解耦。
        tmp_out = os.path.join(tempfile.gettempdir(),
                               f"word_rendered_{os.getpid()}.docx")
        tpl.save(tmp_out)
        doc = Document(tmp_out)
        try:
            os.remove(tmp_out)
        except OSError:
            pass

        _postprocess(doc, tables_map, images_map)
        _cleanup_empty_tags(doc)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        url_path = "/" + output_path.replace("\\", "/").lstrip("/")
        print(json.dumps({"success": True, "outputPath": url_path},
                         ensure_ascii=False))
    except Exception as e:
        safe_err = str(e).encode('utf-8', errors='replace').decode('utf-8')
        print(json.dumps({"success": False, "error": safe_err},
                         ensure_ascii=False))
        sys.exit(1)
    finally:
        for f in temp_files:
            try:
                os.remove(f)
            except OSError:
                pass
