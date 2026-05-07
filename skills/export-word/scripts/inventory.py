#!/usr/bin/env python3
"""
DOCX Template Inventory — 提取模板中所有段落、表格、图片的结构化清单
输出 JSON 格式仓库，包含样式、文本、对齐等信息。
基于 Anthropic docx skill 的 document.py 思想。
"""

import sys
import os
import json
from docx import Document


_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(_SCRIPT_DIR))


def extract_structure(docx_path, output_path=None):
    doc = Document(docx_path)
    structure = {
        "sections": [],
        "paragraphs": [],
        "tables": [],
    }

    # 节信息
    for sec_idx, section in enumerate(doc.sections):
        sec_info = {
            "index": sec_idx,
            "page_width": round(section.page_width / 360000, 2) if section.page_width else None,
            "page_height": round(section.page_height / 360000, 2) if section.page_height else None,
            "header_distance": round(section.header_distance / 360000, 2) if section.header_distance else None,
            "footer_distance": round(section.footer_distance / 360000, 2) if section.footer_distance else None,
        }
        structure["sections"].append(sec_info)

    # 段落
    for p_idx, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        p_info = {
            "index": p_idx,
            "text": para.text[:200],
            "style": para.style.name if para.style else "Normal",
        }
        if para.alignment is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            am = {
                WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
                WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
                WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
            }
            p_info["alignment"] = am.get(para.alignment, str(para.alignment))

        if para.paragraph_format.line_spacing:
            p_info["line_spacing"] = para.paragraph_format.line_spacing

        for run in para.runs:
            r_info = {}
            if run.font.size:
                r_info["font_size"] = run.font.size / 12700
            if run.bold:
                r_info["bold"] = True
            if run.italic:
                r_info["italic"] = True
            if run.font.color and run.font.color.rgb:
                r_info["color"] = str(run.font.color.rgb)
            if r_info:
                p_info["run_style"] = r_info
                break

        structure["paragraphs"].append(p_info)

    # 表格
    for t_idx, table in enumerate(doc.tables):
        t_info = {
            "index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "header_row": [],
        }
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                t_info["header_row"].append(cell.text[:40])
        structure["tables"].append(t_info)

    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(structure, f, ensure_ascii=False, indent=2)
        print(f"已提取结构信息 -> {output_path}")
    else:
        print(json.dumps(structure, ensure_ascii=False, indent=2))

    return structure


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python inventory.py <template.docx> [output.json]", file=sys.stderr)
        sys.exit(1)
    extract_structure(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
