"""
数据可视化构建器
在 Word 文档中嵌入 PNG 图表图片，带标题和来源标注。
"""

import os
from datetime import datetime
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class VisualizationBuilder:
    """数据可视化构建器"""

    def __init__(self, template):
        self.template = template
        self.scheme = template.scheme

    def build(self, doc, chart_paths):
        if not chart_paths:
            return

        tpl = self.template

        tpl._heading(doc, 2, None, "数据可视化分析")

        tpl._body_para(doc,
                       "为直观呈现数据特征与变化趋势，以下图表对各关键指标进行了可视化展示。"
                       "图表采用统一配色方案，便于横向对比分析。")

        for i, path in enumerate(chart_paths):
            if os.path.exists(path):
                p = doc.add_paragraph()
                tpl._para_format(p, before=14, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
                tpl._run(p, f"图 {i + 1}  关键指标可视化", size=10, bold=True, color=self.scheme["primary"])

                try:
                    doc.add_picture(path, width=Inches(5.6))
                except Exception:
                    tpl._body_para(doc, f"[图表文件加载失败: {path}]", first_indent=False)
                    continue

                p2 = doc.add_paragraph()
                tpl._para_format(p2, before=2, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
                from shared.utils import format_datetime_cn
                tpl._run(p2,
                         f"数据来源：实时数据库查询  |  生成时间：{format_datetime_cn()}",
                         size=8, italic=True, color=self.scheme["text_muted"])

        tpl._divider(doc, self.scheme["accent"])
