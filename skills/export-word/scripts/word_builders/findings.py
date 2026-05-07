"""
关键发现构建器
生成带编号的发现与建议列表，中国红强调关键词。
"""

from docx.shared import Pt


class FindingsBuilder:
    """关键发现构建器"""

    def __init__(self, template):
        self.template = template
        self.scheme = template.scheme

    def build(self, doc, findings_list):
        tpl = self.template
        max_items = tpl.config.get("limits", "word_max_findings")

        tpl._heading(doc, 2, None, "关键发现与建议")

        if not findings_list:
            tpl._body_para(doc, "本次分析未识别出显著的数据特征或异常模式。")
            return

        tpl._body_para(doc, "基于上述统计分析及可视化结果，提炼出以下关键发现及对应建议：")

        for i, f in enumerate(findings_list[:max_items], 1):
            p = doc.add_paragraph()
            tpl._para_format(p, before=6, after=3, spacing=1.4)
            tpl._run(p, f"发现 {i}：", size=11, bold=True, color=self.scheme["accent"])

            if isinstance(f, dict):
                content = f.get("finding", f.get("content", str(f)))
                action = f.get("action", f.get("suggestion", ""))
            else:
                content = str(f)
                action = ""

            tpl._run(p, content, size=11, color=self.scheme["text_primary"])

            if action:
                q = doc.add_paragraph()
                tpl._para_format(q, before=2, after=4, spacing=1.3, indent_first=tpl.config.get("layout", "word", "first_indent_cm"))
                tpl._run(q, f"└ 建议：", size=10, bold=True, color=self.scheme["gold"])
                tpl._run(q, action, size=10, color=self.scheme["text_secondary"])

        doc.add_paragraph()
        tpl._divider(doc, self.scheme["accent"])
