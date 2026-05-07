"""
页眉页脚构建器
为所有节设置统一的页眉、页脚和自动页码。
"""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


class HeaderFooterBuilder:
    """页眉页脚构建器"""

    def __init__(self, template):
        self.template = template
        self.scheme = template.scheme

    def build(self, doc, title):
        tpl = self.template

        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            tpl._run(hp, f"{tpl.config.brand_name} · {title}", size=8, color=self.scheme["text_muted"])

            border = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="2" w:space="1" w:color="{self.scheme["border"]}"/>'
                f'</w:pBdr>'
            )
            hp._p.get_or_add_pPr().append(border)

            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            tpl._run(fp, "— 第 ", size=8, color=self.scheme["text_muted"])

            run = fp.add_run()
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            run._r.append(fldChar1)
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = " PAGE "
            run._r.append(instrText)
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run._r.append(fldChar2)

            tpl._run(fp, " 页 —", size=8, color=self.scheme["text_muted"])
