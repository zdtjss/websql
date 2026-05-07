"""
中国商务报告模板
藏蓝色主调、中国红强调、金色点缀，符合国内企业正式报告规范。
"""

from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import WordTemplate


class ChineseReportTemplate(WordTemplate):
    name = "chinese_report"
    label = "中国商务报告"

    def build_cover(self, doc, title, subtitle, dept=""):
        from shared.utils import format_date_cn, generate_report_id

        section = doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        p = doc.add_paragraph()
        self._para_format(p, before=0, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
        self._run(p, f"{self.config.brand_name} · {self.config.brand_tagline}", size=9, color=self.scheme["text_muted"])

        for _ in range(5):
            doc.add_paragraph()

        self._divider(doc, self.scheme["accent"])

        p = doc.add_paragraph()
        self._para_format(p, before=24, after=8, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.1)
        self._run(p, title, size=30, bold=True, color=self.scheme["primary_dark"], font=self.font_en)
        p2 = doc.add_paragraph()
        self._para_format(p2, before=0, after=18, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.1)
        self._run(p2, title, size=30, bold=True, color=self.scheme["primary_dark"])

        self._divider(doc, self.scheme["accent"])

        p3 = doc.add_paragraph()
        self._para_format(p3, before=20, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._run(p3, subtitle, size=15, color=self.scheme["text_secondary"])

        for _ in range(6):
            doc.add_paragraph()

        info_lines = [
            f"编制单位：{dept}" if dept else f"编制单位：{self.config.brand_name} 数据分析中心",
            f"编制日期：{format_date_cn()}",
            f"报告编号：{generate_report_id()}",
            "密级：内部资料",
        ]
        for line in info_lines:
            p = doc.add_paragraph()
            self._para_format(p, before=1, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._run(p, line, size=10, color=self.scheme["text_muted"])

        doc.add_page_break()
