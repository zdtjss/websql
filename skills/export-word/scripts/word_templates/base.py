"""
Word 基础模板
定义所有 Word 报告模板的通用接口和共享排版逻辑。
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


class WordTemplate:
    name = "base"
    label = "基础模板"

    def __init__(self, scheme=None):
        from shared.colors import ColorPalette
        from shared.fonts import FontManager
        from shared.config import SkillConfig

        self.config = SkillConfig()
        self.scheme = scheme or ColorPalette.get("chinese_business")
        self.fonts = FontManager

        self.font_cn = FontManager.get_cn_font()
        self.font_en = FontManager.get_en_font()

    def _rgb(self, hex_color):
        h = hex_color.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def _cell_shading(self, cell, hex_color):
        color = hex_color.lstrip("#")
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _para_format(self, p, before=0, after=0, spacing=1.25, align=None, indent_first=None):
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = spacing
        if align is not None:
            p.alignment = align
        if indent_first is not None:
            pf.first_line_indent = Cm(indent_first)

    def _run(self, p, text, size=11, bold=False, color=None, italic=False, font=None):
        from shared.utils import clean_surrogates
        r = p.add_run()
        fn = font or self.font_cn
        r.font.name = fn
        r._element.rPr.rFonts.set(qn("w:eastAsia"), fn)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = self._rgb(color) if isinstance(color, str) else color
        r.text = clean_surrogates(text)
        return r

    def _divider(self, doc, color=None):
        c = color or self.scheme["divider"]
        p = doc.add_paragraph()
        self._para_format(p, before=4, after=4)
        border = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{c}"/>'
            f'</w:pBdr>'
        )
        p._p.get_or_add_pPr().append(border)

    def _heading(self, doc, level, num, text):
        from shared.utils import num_to_cn, clean_surrogates

        if level == 1:
            p = doc.add_paragraph()
            self._para_format(p, before=28, after=10, spacing=1.2)
            self._run(p, "█", size=18, bold=True, color=self.scheme["accent"])
            self._run(p, "  ", size=8)
            if num:
                self._run(p, f"第{num_to_cn(num)}章  ", size=22, bold=True, color=self.scheme["primary"])
            self._run(p, clean_surrogates(text), size=22, bold=True, color=self.scheme["primary_dark"])
            self._divider(doc, self.scheme["accent"])
        elif level == 2:
            p = doc.add_paragraph()
            self._para_format(p, before=20, after=6, spacing=1.2)
            self._run(p, "┃ ", size=16, bold=True, color=self.scheme["primary"])
            self._run(p, clean_surrogates(text), size=16, bold=True, color=self.scheme["primary"])
        else:
            p = doc.add_paragraph()
            self._para_format(p, before=14, after=4, spacing=1.2)
            self._run(p, "│ ", size=14, bold=True, color=self.scheme["text_secondary"])
            self._run(p, clean_surrogates(text), size=14, bold=True, color=self.scheme["text_primary"])

    def _body_para(self, doc, text, first_indent=True):
        from shared.utils import clean_surrogates
        p = doc.add_paragraph()
        indent = self.config.get("layout", "word", "first_indent_cm", default=0.74) if first_indent else 0
        spacing = self.config.get("layout", "word", "body_line_spacing", default=1.6)
        self._para_format(p, before=2, after=4, spacing=spacing, indent_first=indent)
        self._run(p, clean_surrogates(text), size=11, color=self.scheme["text_primary"])
        return p

    def _create_table(self, doc, headers, rows_data, col_widths=None, align_center=True):
        tbl = doc.add_table(rows=1, cols=len(headers))
        if align_center:
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._run(p, h, size=9, bold=True, color="#FFFFFF")
            self._cell_shading(c, self.scheme["table_header"])

        for j, row_vals in enumerate(rows_data):
            row = tbl.add_row()
            bg = self.scheme["bg_white"] if j % 2 == 0 else self.scheme["table_stripe_word"]
            for k, v in enumerate(row_vals):
                c = row.cells[k]
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._run(p, str(v), size=9, bold=(k == 0))
                self._cell_shading(c, bg)

        return tbl

    def create_document(self):
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = self.font_cn
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_cn)

        section = doc.sections[0]
        section.top_margin = Cm(self.config.get("layout", "word", "page_margin_top_cm"))
        section.bottom_margin = Cm(self.config.get("layout", "word", "page_margin_bottom_cm"))
        section.left_margin = Cm(self.config.get("layout", "word", "page_margin_left_cm"))
        section.right_margin = Cm(self.config.get("layout", "word", "page_margin_right_cm"))

        return doc
